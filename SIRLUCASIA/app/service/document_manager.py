import os
import shutil
import re 
from datetime import datetime

from app.database.document_database import DocumentDatabase
from app.document.document_factory import DocumentFactory
from app.core.action_result import ActionResult
from app.core.action_status import ActionStatus


# Nombres permitidos: letras unicode (acentos y ñ incluidos), dígitos,
# guiones bajos (\w), espacios y guiones. Los separadores de ruta y los
# puntos quedan fuera, así que ".." o "a/b" nunca pasan la validación.
NAME_PATTERN = re.compile(r"[\w -]+", re.UNICODE)


def _is_valid_name(name):
    """True si `name` es un nombre de documento seguro (sin rutas)."""
    if not name or not name.strip():
        return False
    if any(sep in name for sep in ("/", "\\", "..")):
        return False
    return bool(NAME_PATTERN.fullmatch(name))


# ==================================================
# DocumentManager
# Controla la creación y administración de documentos
# ==================================================

class DocumentManager:

    def __init__(self):
        self.database = DocumentDatabase()
        self.path = "documents"
        os.makedirs(self.path, exist_ok=True)

        print("=" * 50)
        print("[DocumentManager]")
        print("Inicializado correctamente.")
        print("=" * 50)

    # ==================================================
    # Traductor de formatos
    # ==================================================
    def _normalize_format(self, format_name):
        if not format_name:
            return "docx"

        format_name = format_name.lower().strip()
        aliases = {
            "documento": "docx", "word": "docx", "doc": "docx", "docx": "docx",
            "txt": "txt", "texto": "txt", "nota": "txt", "nota de texto": "txt",
            "bloc": "txt", "bloc de notas": "txt", "notepad": "txt",
            "pdf": "pdf",
            "excel": "xlsx", "xlsx": "xlsx", "hoja": "xlsx", "hoja de calculo": "xlsx",
            "powerpoint": "pptx", "power point": "pptx", "ppt": "pptx", "pptx": "pptx",
            "presentacion": "pptx", "presentación": "pptx",
            "json": "json",
            "markdown": "md", "md": "md"
        }
        # CAMBIO: si no está en aliases, devuelve el mismo valor
        return aliases.get(format_name, format_name)


    # ==================================================
    # Router
    # ==================================================
    def execute(self, data):
        command = data.get("command")
        method = getattr(self, command, None)

        if method is None:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command=command,
                message=f"No existe la acción '{command}'."
            )

        return method(data)
    # ==================================================
    # Crear documento
    # ==================================================
    # ==================================================

    def create(self, data):
        name = data.get("topic")
        if not name:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="create",
                message="No especificaste el nombre del documento."
            )

        # Validar nombre
        if not _is_valid_name(name):
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="create",
                message="El nombre del documento es inválido."
            )

        # Validar contenido
        content = data.get("content", "")
        if not content.strip():
            # CAMBIO: permitir documento vacío
            content = ""

        # Validar formato
        format_name = self._normalize_format(data.get("format"))
        if not format_name:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="create",
                message="Formato no soportado."
            )

        extension = self.database.find(format_name)
        valid_extensions = [".docx", ".txt", ".pdf", ".xlsx", ".pptx", ".json", ".md"]
        if extension not in valid_extensions:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="create",
                message="Formato no soportado."
            )

        filename = f"{name}{extension}"
        filepath = os.path.join(self.path, filename)

        # Generar nombre alternativo si ya existe
        counter = 1
        while os.path.exists(filepath):
            filename = f"{name}_{counter}{extension}"
            filepath = os.path.join(self.path, filename)
            counter += 1

        # Crear documento
        try:
            DocumentFactory.create(extension, filepath, content)
            return ActionResult(
                success=True,
                status=ActionStatus.SUCCESS,
                module="document",
                command="create",
                message=f"Documento '{filename}' creado correctamente.",
                data={"filename": filename}
            )
        except Exception as e:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="create",
                message=f"No pude crear el documento: {e}"
            )


        # ==================================================
        # Leer documento
        # ==================================================
    def read(self, data):
        name = data.get("topic")
        if not name:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="read",
                message="No especificaste el nombre del documento."
            )

        filepath, extension = self._get_document(name)
        if filepath is None:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="read",
                message="No encontré ese documento."
            )

        try:
            if extension in [".txt", ".md", ".json"]:
                with open(filepath, "r", encoding="utf-8") as file:
                    content = file.read()
            elif extension == ".docx":
                from docx import Document
                doc = Document(filepath)
                content = "\n".join(p.text for p in doc.paragraphs)
            elif extension == ".pdf":
                return ActionResult(
                    success=False,
                    status=ActionStatus.WARNING,
                    module="document",
                    command="read",
                    message="La lectura de PDF llegará en una próxima versión."
                )
            else:
                return ActionResult(
                    success=False,
                    status=ActionStatus.ERROR,
                    module="document",
                    command="read",
                    message="Formato no soportado."
                )

            return ActionResult(
                success=True,
                status=ActionStatus.SUCCESS,
                module="document",
                command="read",
                message="Contenido leído correctamente.",
                data={"content": content}
            )
        except Exception as e:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="read",
                message=str(e)
            )

    # ==================================================
    # Escribir documento
    # ==================================================
    def write(self, data):
        name = data.get("topic")
        content = data.get("content", "")

        if not name:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="write",
                message="No especificaste el nombre del documento."
            )

        filepath, extension = self._get_document(name)
        if filepath is None:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="write",
                message="No encontré ese documento."
            )

        try:
            if extension in [".txt", ".md"]:
                prefix = "\n" if os.path.getsize(filepath) > 0 else ""
                with open(filepath, "a", encoding="utf-8") as file:
                    file.write(prefix + content)
            elif extension == ".docx":
                from docx import Document
                doc = Document(filepath)
                doc.add_paragraph(content)
                doc.save(filepath)
            elif extension == ".json":
                return ActionResult(
                    success=False,
                    status=ActionStatus.WARNING,
                    module="document",
                    command="write",
                    message="Por seguridad todavía no puedo modificar archivos JSON."
                )
            elif extension == ".pdf":
                return ActionResult(
                    success=False,
                    status=ActionStatus.WARNING,
                    module="document",
                    command="write",
                    message="La edición de PDF estará disponible en una próxima versión."
                )

            return ActionResult(
                success=True,
                status=ActionStatus.SUCCESS,
                module="document",
                command="write",
                message=f"Contenido agregado a '{os.path.basename(filepath)}'."
            )
        except Exception as e:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="write",
                message=str(e)
            )

    # ==================================================
    # Eliminar documento
    # ==================================================
    def delete(self, data):
        name = data.get("topic")
        filepath, extension = self._get_document(name)
        if filepath is None:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="delete",
                message="No encontré ese documento."
            )

        try:
            os.remove(filepath)
            return ActionResult(
                success=True,
                status=ActionStatus.SUCCESS,
                module="document",
                command="delete",
                message=f"Documento '{os.path.basename(filepath)}' eliminado."
            )
        except Exception as e:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="delete",
                message=str(e)
            )

    # ==================================================
    # Renombrar documento
    # ==================================================
    def rename(self, data):

        old_name = data.get("old_name") or data.get("topic")
        new_name = data.get("new_name")

        filepath, extension = self._get_document(old_name)

        if filepath is None:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="rename",
                message="No encontré ese documento."
            )

        new_path = os.path.join(
            self.path,
            f"{new_name}{extension}"
        )

        # 🔧 CAMBIO: Generar nombre alternativo si ya existe
        counter = 1
        while os.path.exists(new_path):
            new_path = os.path.join(
                self.path,
                f"{new_name}_{counter}{extension}"
            )
            counter += 1

        try:
            os.rename(filepath, new_path)

            return ActionResult(
                success=True,
                status=ActionStatus.SUCCESS,
                module="document",
                command="rename",
                message=f"'{old_name}' fue renombrado a '{os.path.basename(new_path)}'."
            )

        except Exception as e:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="rename",
                message=str(e)
            )



    # ==================================================
    # Copiar documento
    # ==================================================

    # ==================================================
    # Copiar documento
    # ==================================================
    def copy(self, data):
        old_name = data.get("old_name")
        new_name = data.get("new_name")
        filepath, extension = self._get_document(old_name)
        if filepath is None:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="copy",
                message="No encontré ese documento."
            )

        new_path = os.path.join(self.path, f"{new_name}{extension}")
        try:
            shutil.copy(filepath, new_path)
            return ActionResult(
                success=True,
                status=ActionStatus.SUCCESS,
                module="document",
                command="copy",
                message=f"'{old_name}' copiado como '{new_name}'."
            )
        except Exception as e:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="copy",
                message=str(e)
            )

    # ==================================================
    # Mover documento
    # ==================================================
    def move(self, data):
        name = data.get("topic")
        new_path = data.get("new_path")
        filepath, extension = self._get_document(name)

        if filepath is None:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="move",
                message="No encontré ese documento."
            )

        try:
            shutil.move(filepath, new_path)
            return ActionResult(
                success=True,
                status=ActionStatus.SUCCESS,
                module="document",
                command="move",
                message="Documento movido correctamente."
            )
        except Exception as e:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="move",
                message=str(e)
            )
    # ==================================================
    # Listar documentos
    # ==================================================    
    def list_documents(self, data=None):

        try:
            files = os.listdir(self.path)

            return ActionResult(
                success=True,
                status=ActionStatus.SUCCESS,
                module="document",
                command="list_documents",
                message="Documentos obtenidos.",
                data=files
            )

        except Exception as e:

            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="list_documents",
                message=str(e)
           )

    # ==================================================
    # Información del documento
    # ==================================================
    def info(self, data):
        name = data.get("topic")

        filepath, extension = self._get_document(name)

        if filepath is None:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="info",
                message="No encontré ese documento."
            )

        try:
            stat = os.stat(filepath)

            return ActionResult(
                success=True,
                status=ActionStatus.SUCCESS,
                module="document",
                command="info",
                message="Información obtenida.",
                data={
                    "name": os.path.basename(filepath),
                    "extension": extension,
                    "size": stat.st_size,
                    "created": datetime.fromtimestamp(stat.st_ctime),
                    "modified": datetime.fromtimestamp(stat.st_mtime)
                }
            )

        except Exception as e:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="info",
                message=str(e)
            )

    # ==================================================
    # Buscar documento
    # ==================================================
    def search(self, data):
        name = data.get("topic")

        filepath, extension = self._get_document(name)

        if filepath is None:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="search",
                message="No encontré ese documento."
            )

        return ActionResult(
            success=True,
            status=ActionStatus.SUCCESS,
            module="document",
            command="search",
            message="Documento encontrado.",
            data={
                "path": filepath,
                "extension": extension
            }
        )

    # ==================================================
    # Última modificación
    # ==================================================
    def modified(self, data):
        name = data.get("topic")

        filepath, extension = self._get_document(name)

        if filepath is None:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="modified",
                message="No encontré ese documento."
            )

        try:
            modified = datetime.fromtimestamp(
                os.path.getmtime(filepath)
            )

            return ActionResult(
                success=True,
                status=ActionStatus.SUCCESS,
                module="document",
                command="modified",
                message="Fecha obtenida.",
                data={
                    "modified": modified
                    }
            )

        except Exception as e:
            return ActionResult(
                success=False,
                status=ActionStatus.ERROR,
                module="document",
                command="modified",
                message=str(e)
            )

    # ==================================================
    # Buscar documento físicamente
    # ==================================================
    def find_document(self, name):
        if not name:
            return None, None
        if not os.path.exists(self.path):
            return None, None

        for file in os.listdir(self.path):
            filename, extension = os.path.splitext(file)
            # Comparar tanto con nombre sin extensión como con nombre completo
            if filename.lower() == name.lower() or file.lower() == name.lower():
                return os.path.join(self.path, file), extension
        return None, None

    # ==================================================
    # Obtener documento
    # ==================================================
    def _get_document(self, name):
        return self.find_document(name)

    # ==================================================
    # Verificar existencia
    # ==================================================
    def exists(self, data):

        if isinstance(data, dict):
            name = data.get("topic")
        else:
             name = data

        filepath, _ = self._get_document(name)

        return filepath is not None